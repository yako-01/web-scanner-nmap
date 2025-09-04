from operator import add
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# Diccionario con los valores de prueba
replacements = {
    "ENTREPRISE": "AAAAAA",
    "DATE": datetime.datetime.now().strftime("%d/%m/%Y"),
    "SITE": "testphp.vulnweb.com",
    "IP_ADDRESS": "192.168.1.10"
}



def allow_table_row_break_across_pages(table):
    for row in table.rows:
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        tblHeader = OxmlElement('w:cantSplit')
        trPr.append(tblHeader)


def replace_text_in_paragraph(paragraph, replacements):
    for key, val in replacements.items():
        tag = f"*{key}*"
        if tag in paragraph.text:
            for run in paragraph.runs:
                if tag in run.text:
                    run.text = run.text.replace(tag, str(val))

def replace_tags_in_docx(template_path, output_path, replacements):
    doc = Document(template_path)

    def replace_text_in_table(table):
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_text_in_paragraph(paragraph, replacements)

    # Reemplazo en párrafos
    for paragraph in doc.paragraphs:
        replace_text_in_paragraph(paragraph, replacements)

    # Reemplazo en tablas
    for table in doc.tables:
        replace_text_in_table(table)

    # Reemplazo en cabeceras y pies
    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            replace_text_in_paragraph(paragraph, replacements)
        for paragraph in section.footer.paragraphs:
            replace_text_in_paragraph(paragraph, replacements)

    doc.save(output_path)

# Prueba con tu template
#replace_tags_in_docx("rapport_template.docx", "resultado3.docx", replacements)
for table in Document("resultado3.docx").tables:
    allow_table_row_break_across_pages(table)
