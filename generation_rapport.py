import os
import json
import sys
from typing import Dict, List, Tuple
from docx import Document
import datetime



INTERMEDIATE_DIR = "intermediate"
PORTS = [21,22,23,25,53,80,110,143,443,445,465,993,995,1433,1521,3306,3389,5432,5900,8080,8443,27017]
PORT_SERVICE_MAP = {
    "21": "FTP",
    "22": "SSH",
    "23": "Telnet",
    "25": "SMTP",
    "53": "DNS",
    "80": "HTTP",
    "110": "POP3",
    "143": "IMAP",
    "443": "HTTPS",
    "445": "SMB",
    "465": "SMTPS",
    "993": "IMAPS",
    "995": "POP3S",
    "1433": "MSSQL",
    "1521": "Oracle DB",
    "3306": "MySQL",
    "3389": "RDP",
    "5432": "PostgreSQL",
    "5900": "VNC",
    "8080": "HTTP Proxy",
    "8443": "HTTPS Proxy",
    "27017": "MongoDB",
}

def completar_closed_ports(intermediate_file, ports_list):
    """
    :param intermediate_file: Ruta al fichero intermedio generado por el script bash
    :param ports_list: Lista de puertos definidos en scan_nmap.sh
    :return: diccionario con todas las secciones (incluyendo CLOSED_PORTS completado)
    """
    data = {}
    current_key = None

    # Leer fichero intermedio
    with open(intermediate_file, "r") as f:
        for line in f:
            line = line.strip()
            if not line:
                continue
            if line.endswith(":"):
                current_key = line[:-1]
                data[current_key] = []
            else:
                if current_key:
                    data[current_key].append(line)

    # Pasar puertos abiertos a int
    open_ports = [int(p) for p in data.get("OPEN_PORTS", []) if p.isdigit()]

    #filtered_ports = [int(p) for p in data.get("FILTERED_PORTS", []) if p.isdigit()]

    # Calcular cerrados como diferencia
    closed_ports = [str(p) for p in ports_list if p not in open_ports]

    # Rellenar en data
    data["CLOSED_PORTS"] = closed_ports

    return data


def get_service_name(port, fallback="desconocido"):
    """Devuelve el nombre amigable del servicio asociado a un puerto"""
    return PORT_SERVICE_MAP.get(port, fallback)


def remplir_tableau_ports(docx_template, output_docx, data):
    doc = Document(docx_template)

    # Recorrer tabla(s) del doc
    for table in doc.tables:
        # Detectamos si la tabla es la de puertos (col cabecera contiene "Numéro du port")
        headers = [cell.text.strip().lower() for cell in table.rows[0].cells]
        if "numéro du port" in headers[0]:
            # Borrar las filas excepto la cabecera
            for row in table.rows[1:]:
                tbl = table._tbl
                tbl.remove(row._tr)

            # Añadir filas de puertos abiertos
            for port in data.get("OPEN_PORTS", []):
                row_cells = table.add_row().cells
                row_cells[0].text = port
                row_cells[1].text = get_service_name(port)
                row_cells[2].text = "Ouvert"
            
            """for port in data.get("FILTERED_PORTS", []):
                row_cells = table.add_row().cells
                row_cells[0].text = port
                row_cells[1].text = get_service_name(port)
                row_cells[2].text = "Filtré"
            """
            # Añadir filas de puertos cerrados
            for port in data.get("CLOSED_PORTS", []):
                row_cells = table.add_row().cells
                row_cells[0].text = port
                row_cells[1].text = get_service_name(port)
                row_cells[2].text = "Fermé"

    doc.save(output_docx)


def extract_scan_id(filename: str) -> str:
    """
    Extrae el scan_id (timestamp) del nombre de archivo intermedio.
    Ejemplo: nmap_testphp.vulnweb.com_01092025_103149.txt -> 01092025_103149
    """
    base = os.path.basename(filename)  # nmap_testphp.vulnweb.com_01092025_103149.txt
    parts = base.split("_")
    if len(parts) >= 3:
        return parts[-1].replace(".txt", "")  # -> 01092025_103149
    else:
        raise ValueError(f"No se pudo extraer scan_id del archivo {filename}")

"""
resultado = completar_closed_ports("intermediate/nmap_testphp.vulnweb.com_01092025_103149.txt", PORTS)
print(resultado)

output_file = f"./reports/informe_generado{datetime.datetime.now().strftime('%d%m%Y_%H%M%S')}.docx"
remplir_tableau_ports("./rapport_template.docx", output_file, resultado)
"""



def tags_valors(intermediate_file):
    target = ""
    with open(intermediate_file, "r", encoding="utf-8", errors="ignore") as f:
        for line in f:
            line = line.strip()
            if not line:
                continue
            if line.startswith("TARGET:"):
                target = line.split(":", 1)[1].strip()
                break

    info = {
        "ENTREPRISE": "aaaa",
        "DATE": datetime.datetime.now().strftime("%d/%m/%Y"),
        "SITE": target
    }
    return info


def find_latest_zap_report() -> str | None:
    reports_dir = os.path.join("reports")
    if not os.path.isdir(reports_dir):
        return None
    candidates: List[Tuple[float, str]] = []
    for name in os.listdir(reports_dir):
        if not name.endswith(".txt"):
            continue
        if name.startswith("rapport_zap_"):
            full = os.path.join(reports_dir, name)
            try:
                mtime = os.path.getmtime(full)
            except OSError:
                continue
            candidates.append((mtime, full))
    if not candidates:
        return None
    candidates.sort(key=lambda x: x[0], reverse=True)
    return candidates[0][1]


def replace_tags_in_docx(template_path: str, output_path: str, replacements: dict):
    """
    Reemplaza en el docx todos los *TAGS* por los valores dados en replacements.

    :param template_path: Ruta al documento template (docx)
    :param output_path: Ruta al documento de salida (docx)
    :param replacements: Diccionario con pares { "TAG": "valor" }
                         Ejemplo: { "DATE": "01/09/2025", "SITE": "testphp.vulnweb.com" }
    """
    doc = Document(template_path)

    def replace_text_in_paragraph(paragraph, replacements):
        for key, val in replacements.items():
            tag = f"*{key}*"
            if tag in paragraph.text:
                # Recorrer runs porque los tags pueden estar fragmentados
                for run in paragraph.runs:
                    if tag in run.text:
                        run.text = run.text.replace(tag, str(val))

    def replace_text_in_table(table, replacements):
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_text_in_paragraph(paragraph, replacements)

    # Reemplazo en párrafos principales
    for paragraph in doc.paragraphs:
        replace_text_in_paragraph(paragraph, replacements)

    # Reemplazo en cabeceras/pies de página
    for section in doc.sections:
        header = section.header
        for paragraph in header.paragraphs:
            replace_text_in_paragraph(paragraph, replacements)
        footer = section.footer
        for paragraph in footer.paragraphs:
            replace_text_in_paragraph(paragraph, replacements)

    # Guardar documento final
    doc.save(output_path)



def parse_zap_report(report_path: str) -> List[Dict[str, str | int]]:
    """
    Parsea un informe ZAP de texto (formato similar a reports/rapport_zap_example.txt)
    y devuelve una lista de hallazgos con:
      - title (Test)
      - risk (Risque)
      - description (Description)
      - recommendation (Solution)
      - detected_count (Cas détectés - número)
      - cases (texto con la lista de "Cas détectés")
    """
    findings: List[Dict[str, str | int]] = []
    if not report_path or not os.path.exists(report_path):
        return findings

    def flush_current():
        nonlocal current_title, current_risk, current_description_lines, current_solution_lines, current_detected_count, current_cases_lines
        if current_title:
            finding: Dict[str, str | int] = {
                "title": current_title.strip("= "),
                "risk": current_risk.strip() if current_risk else "",
                "description": "\n".join([l.strip() for l in current_description_lines]).strip(),
                "recommendation": "\n".join([l.strip() for l in current_solution_lines]).strip(),
                "detected_count": current_detected_count if current_detected_count is not None else 0,
                "cases": "\n".join(current_cases_lines).strip(),
            }
            findings.append(finding)

    current_title: str = ""
    current_risk: str = ""
    current_description_lines: List[str] = []
    current_solution_lines: List[str] = []
    current_detected_count: int | None = None
    current_cases_lines: List[str] = []
    mode: str | None = None  # 'desc' | 'sol' | 'cases' | None

    with open(report_path, "r", encoding="utf-8", errors="ignore") as fh:
        for raw in fh:
            line = raw.rstrip("\n")
            
            # Nueva sección de vulnerabilidad
            if line.startswith("=== ") and line.endswith(" ==="):
                flush_current()
                current_title = line.strip("= ").strip()
                current_risk = ""
                current_description_lines = []
                current_solution_lines = []
                current_detected_count = None
                current_cases_lines = []
                mode = None
                continue
            
            # Línea de riesgo
            if line.startswith("Risque :"):
                risk_part = line.split(":", 1)[1].split("|")[0].strip()
                current_risk = risk_part
                mode = None
                continue
            
            # Inicio de sección Description
            if line.strip() == "Description :":
                mode = "desc"
                continue
            
            # Inicio de sección Solution
            if line.strip() == "Solution :":
                mode = "sol"
                continue
            
            # Línea de Cas détectés
            if line.strip().startswith("Cas détectés ("):
                try:
                    # Extraer número entre paréntesis: "Cas détectés (44) :"
                    start = line.find("(") + 1
                    end = line.find(")")
                    if start > 0 and end > start:
                        number_str = line[start:end].strip()
                        current_detected_count = int(number_str)
                except Exception:
                    current_detected_count = 0
                mode = "cases"
                current_cases_lines = []
                continue
            
            # Acumular líneas según el modo actual
            if mode == "desc":
                if line.strip() == "":
                    continue
                # Si encontramos cambio de bloque, no acumular más
                if line.strip() == "Solution :":
                    mode = "sol"
                    continue
                if line.strip().startswith("Cas détectés"):
                    mode = "cases"
                    current_cases_lines = []
                    continue
                current_description_lines.append(line)
                
            elif mode == "sol":
                if line.strip() == "":
                    continue
                if line.strip().startswith("Cas détectés"):
                    mode = "cases"
                    current_cases_lines = []
                    continue
                current_solution_lines.append(line)

            elif mode == "cases":
                # Termina cuando llega una nueva sección '==='
                if line.startswith("=== ") and line.endswith(" ==="):
                    # Retrocede una iteración: for se encargará de flush al ver nueva sección
                    # Pero como ya detectamos aquí, hacemos flush y reiniciamos manualmente
                    flush_current()
                    current_title = line.strip("= ").strip()
                    current_risk = ""
                    current_description_lines = []
                    current_solution_lines = []
                    current_detected_count = None
                    current_cases_lines = []
                    mode = None
                    continue
                if line.strip() == "":
                    continue
                current_cases_lines.append(line)

    # Flush final
    flush_current()
    return findings


def _table_headers_lower(table) -> List[str]:
    return [cell.text.strip().lower() for cell in table.rows[0].cells]


def _find_table_after_heading(doc: Document, heading_contains: str, min_columns: int | None = None):
    """
    Busca la tabla inmediatamente después de un párrafo cuyo texto contenga heading_contains.
    """
    body = doc._element.body
    found_heading = False
    for child in body.iterchildren():
        tag = child.tag.lower()
        if tag.endswith('p'):
            # párrafo
            text = "".join([r.text or "" for r in child.iter() if hasattr(r, 'text')])
            if text and heading_contains.lower() in text.strip().lower():
                found_heading = True
        elif tag.endswith('tbl') and found_heading:
            from docx.table import Table
            table = Table(child, doc)
            if min_columns and table.rows and len(table.rows[0].cells) < min_columns:
                return table
            return table
    return None


def fill_zap_vulnerabilities_table(output_docx: str, findings: List[Dict[str, str | int]]):
    """
    Rellena la tabla "Analyse de la configuration du serveur Internet" con columnas:
    Test | Description | Risque | Cas Détectés | Recommendation
    """
    if not findings:
        return
    doc = Document(output_docx)
    # Intento: tabla tras el título indicado
    table = _find_table_after_heading(doc, "La configuration du site internet présente les vulnérabilités suivantes :", min_columns=5)
    # Fallback por cabeceras
    if table is None:
        for t in doc.tables:
            headers = _table_headers_lower(t)
            if len(headers) >= 5 and \
               "test" in headers[0] and \
               "description" in headers[1] and \
               "risque" in headers[2] and \
               "cas" in headers[3] and \
               "recommendation" in headers[4]:
                table = t
                break
    if table is None:
        doc.save(output_docx)
        return
    # Limpiar filas excepto cabecera
    for row in table.rows[1:]:
        table._tbl.remove(row._tr)
    # Añadir hallazgos
    for f in findings:
        row = table.add_row().cells
        row[0].text = str(f.get("title", ""))
        row[1].text = str(f.get("description", ""))
        row[2].text = str(f.get("risk", ""))
        row[3].text = str(f.get("detected_count", 0))
        row[4].text = str(f.get("recommendation", ""))
    doc.save(output_docx)




def fill_zap_configuration_table(output_docx: str, findings: List[Dict[str, str | int]]):
    """
    Rellena la tabla "Configuration du site Internet" con columnas:
    Test | Cas Détectés (lista de URLs/Parámetros)
    """
    if not findings:
        return
    doc = Document(output_docx)
    # Intento: tabla tras el título indicado
    table = _find_table_after_heading(doc, "Les cas détectés sont les suivants :", min_columns=2)
    # Fallback por cabeceras
    if table is None:
        for t in doc.tables:
            headers = _table_headers_lower(t)
            if len(headers) >= 2 and "test" in headers[0] and "cas" in headers[1]:
                table = t
                break
    if table is None:
        doc.save(output_docx)
        return
    # Limpiar
    for row in table.rows[1:]:
        table._tbl.remove(row._tr)
    # Rellenar con lista de casos por test
    for f in findings:
        row = table.add_row().cells
        row[0].text = str(f.get("title", ""))
        cases_text = str(f.get("cases", "")).strip()
        if not cases_text and f.get("detected_count", 0):
            cases_text = f"{f.get('detected_count')} cas détectés"
        row[1].text = cases_text
    doc.save(output_docx)


def remplir_tableaux_zap(output_docx: str, zap_report_path: str | None = None):
    """
    Carga hallazgos ZAP y rellena ambas tablas en un DOCX ya existente.
    Si no se pasa ruta, intenta usar el último reporte ZAP o el de ejemplo.
    """
    if not zap_report_path:
        zap_report_path = find_latest_zap_report()
    if not zap_report_path or not os.path.exists(zap_report_path):
        print(f"ERROR: No se encontró el reporte ZAP en {zap_report_path}")
        return
    
    print(f"DEBUG: Parseando reporte ZAP desde {zap_report_path}")
    findings = parse_zap_report(zap_report_path)
    print(f"DEBUG: Se encontraron {len(findings)} hallazgos")
    
    if not findings:
        print("ERROR: No se pudieron extraer hallazgos del reporte ZAP")
        return
    
    # Debug: mostrar el primer hallazgo
    if findings:
        first = findings[0]
        print(f"DEBUG Primer hallazgo:")
        print(f"  Title: {first.get('title', 'N/A')}")
        print(f"  Risk: {first.get('risk', 'N/A')}")
        print(f"  Description length: {len(str(first.get('description', '')))}")
        print(f"  Recommendation length: {len(str(first.get('recommendation', '')))}")
        print(f"  Detected count: {first.get('detected_count', 'N/A')}")
    
    fill_zap_vulnerabilities_table(output_docx, findings)
    fill_zap_configuration_table(output_docx, findings)
    print(f"DEBUG: Tablas ZAP rellenadas en {output_docx}")




def replace_text_in_paragraph(paragraph, replacements):
    for key, val in replacements.items():
        tag = f"*{key}*"
        if tag in paragraph.text:
            for run in paragraph.runs:
                if tag in run.text:
                    run.text = run.text.replace(tag, str(val))

def replace_tags_in_docx(template_path, output_path, replacements):
    doc = Document(template_path)

    def replace_text_in_table(table):
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_text_in_paragraph(paragraph, replacements)

    # Reemplazo en párrafos
    for paragraph in doc.paragraphs:
        replace_text_in_paragraph(paragraph, replacements)

    # Reemplazo en tablas
    for table in doc.tables:
        replace_text_in_table(table)

    # Reemplazo en cabeceras y pies
    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            replace_text_in_paragraph(paragraph, replacements)
        for paragraph in section.footer.paragraphs:
            replace_text_in_paragraph(paragraph, replacements)

    doc.save(output_path)



#remplir_tableaux_zap("./reports/informe_zap.docx", "./reports/rapport_zap_example.txt")
    
def generar_informe(intermediate_file):
    scan_id = extract_scan_id(intermediate_file)
    template_path = "./rapport_template2.docx"
    output_file = f"./reports/informe_definitivo_{scan_id}.docx"
    
    datos_nmap = completar_closed_ports(intermediate_file, PORTS)
    remplir_tableau_ports(template_path, output_file, datos_nmap)
    replace_tags_in_docx(output_file, output_file, tags_valors(intermediate_file))


    
    return output_file  # ⬅ devolvemos el archivo generado

if __name__ == "__main__":
    intermediate_file = sys.argv[1]
    output_file = generar_informe(intermediate_file)
    print(f"Informe generado en: {output_file}")


    with open("./intermediate/ultimo_output.txt", "w") as f:
        f.write(output_file)

    print(f"Reporte generado: {output_file}")

    """
    # Uso: python generation_rapport.py intermediate/nmap_...txt [ruta_reporte_zap_txt]
    if len(sys.argv) < 2:
        print("Uso: python generation_rapport.py <ruta_intermedio_nmap> [ruta_reporte_zap]")
        sys.exit(1)

    intermediate_file = sys.argv[1]
    zap_report_arg = sys.argv[2] if len(sys.argv) >= 3 else None

    # 1) Extraer scan_id del fichero intermedio Nmap para nombrar el DOCX final
    scan_id = extract_scan_id(intermediate_file)

    # 2) Construir diccionario de tags
    tags = tags_valors(intermediate_file)

    # 3) Preparar rutas
    template_path = "./rapport_template2.docx"
    output_file = f"./reports/informe_definitivo_{scan_id}.docx"

    # 4) Reemplazar tags en un nuevo DOCX (no modificamos ninguno previo)
    #replace_tags_in_docx(template_path, output_file, tags)

    # 5) Rellenar tabla de puertos Nmap
    datos_nmap = completar_closed_ports(intermediate_file, PORTS)
    # Para Nmap, la función espera leer desde template; generaremos sobre el DOCX ya creado
    # Creamos una versión temporal para reutilizar la función existente
    remplir_tableau_ports(template_path, output_file, datos_nmap)

    # 6) Rellenar tablas de ZAP (vulns y configuración)
    remplir_tableaux_zap(output_file, zap_report_arg)

    print(f"Documento generado: {output_file}")
    """




"""
def find_latest_intermediate_file(target: str | None = None) -> str | None:
    if not os.path.isdir(INTERMEDIATE_DIR):
        return None
    candidates: List[Tuple[float, str]] = []
    for name in os.listdir(INTERMEDIATE_DIR):
        if not name.startswith("nmap_") or not name.endswith(".txt"):
            continue
        if target and (f"nmap_{target}_" not in name):
            continue
        full = os.path.join(INTERMEDIATE_DIR, name)
        try:
            mtime = os.path.getmtime(full)
        except OSError:
            continue
        candidates.append((mtime, full))
    if not candidates:
        return None
    candidates.sort(key=lambda x: x[0], reverse=True)
    return candidates[0][1]


def parse_intermediate_file(path: str) -> Dict[str, List[str] | str]:
    sections: Dict[str, List[str] | str] = {
        "TARGET": "",
        "TIMESTAMP": "",
        "RESOLVED_IPS": [],
        "OPEN_PORTS": [],
        "CLOSED_PORTS": [],
        "FILTERED_PORTS": [],
        "SERVICE_VERSIONS": [],
    }
    if not path or not os.path.exists(path):
        return sections

    current: str | None = None
    with open(path, "r", encoding="utf-8", errors="ignore") as fh:
        for raw in fh:
            line = raw.rstrip("\n")
            if line.startswith("TARGET:"):
                sections["TARGET"] = line.split(":", 1)[1].strip()
                current = None
                continue
            if line.startswith("TIMESTAMP:"):
                sections["TIMESTAMP"] = line.split(":", 1)[1].strip()
                current = None
                continue
            if line in ("RESOLVED_IPS:", "OPEN_PORTS:", "CLOSED_PORTS:", "FILTERED_PORTS:", "SERVICE_VERSIONS:"):
                current = line[:-1]
                continue
            if current:
                if line.strip() == "":
                    continue
                # Append raw line content to the active list
                lst: List[str] = sections[current]  # type: ignore[assignment]
                lst.append(line.strip())
    return sections


def load_data(target: str | None = None) -> Dict[str, List[str] | str]:
    path = find_latest_intermediate_file(target)
    return parse_intermediate_file(path) if path else {
        "TARGET": target or "",
        "TIMESTAMP": "",
        "RESOLVED_IPS": [],
        "OPEN_PORTS": [],
        "CLOSED_PORTS": [],
        "FILTERED_PORTS": [],
        "SERVICE_VERSIONS": [],
    }


def as_json(target: str | None = None) -> str:
    return json.dumps(load_data(target), ensure_ascii=False, indent=2)


if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(description="Leer datos intermedios de Nmap y mostrarlos en JSON")
    parser.add_argument("target", nargs="?", help="IP o dominio para filtrar el fichero intermedio más reciente")
    parser.add_argument("--json", action="store_true", help="Imprimir en formato JSON")
    args = parser.parse_args()

    data = load_data(args.target)
    if args.json:
        print(json.dumps(data, ensure_ascii=False, indent=2))
    else:
        print(f"TARGET: {data.get('TARGET','')}")
        print(f"TIMESTAMP: {data.get('TIMESTAMP','')}")
        print("RESOLVED_IPS:")
        for v in data.get("RESOLVED_IPS", []):
            print(v)
        print("OPEN_PORTS:")
        for v in data.get("OPEN_PORTS", []):
            print(v)
        print("CLOSED_PORTS:")
        for v in data.get("CLOSED_PORTS", []):
            print(v)
        print("FILTERED_PORTS:")
        for v in data.get("FILTERED_PORTS", []):
            print(v)
        print("SERVICE_VERSIONS:")
        for v in data.get("SERVICE_VERSIONS", []):
            print(v)

"""
